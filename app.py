"""
PredictiveAI - Streamlit Application
Predictive Maintenance Scheduling with Multi-LLM Support
"""

import os
import sys
import io
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import streamlit as st
from datetime import datetime

# Add current directory for imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from pattern_detector import (
    detect_degradation_patterns,
    detect_near_misses,
    generate_maintenance_schedule,
    get_equipment_health_scores,
    get_trend_data,
)
from ai_engine import analyze_logs_with_llm, chat_query
from ml_features import (
    detect_anomalies_isolation_forest,
    predict_failure_probability,
    generate_report_data,
)

# ── Page Config ──────────────────────────────────────────────────
st.set_page_config(
    page_title="PredictiveAI - Maintenance Intelligence",
    page_icon="⚙️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS for Light Theme ───────────────────────────────────
st.markdown(
    """
<style>
    /* Global light theme overrides */
    .stApp {
        background-color: #FFFFFF;
    }

    /* Header styling */
    .main-header {
        background: linear-gradient(135deg, #2563EB 0%, #7C3AED 100%);
        padding: 1.5rem 2rem;
        border-radius: 12px;
        margin-bottom: 1.5rem;
        color: white;
    }
    .main-header h1 {
        color: white !important;
        margin: 0;
        font-size: 1.8rem;
    }
    .main-header p {
        color: rgba(255,255,255,0.85);
        margin: 0.3rem 0 0 0;
        font-size: 0.95rem;
    }

    /* Metric cards */
    .metric-card {
        background: #F8FAFC;
        border: 1px solid #E2E8F0;
        border-radius: 10px;
        padding: 1.2rem;
        text-align: center;
        transition: box-shadow 0.2s;
    }
    .metric-card:hover {
        box-shadow: 0 4px 12px rgba(0,0,0,0.08);
    }
    .metric-value {
        font-size: 2rem;
        font-weight: 700;
        color: #1E293B;
    }
    .metric-label {
        font-size: 0.85rem;
        color: #64748B;
        margin-top: 0.3rem;
    }

    /* Status badges */
    .badge-critical { background: #FEE2E2; color: #DC2626; padding: 3px 10px; border-radius: 12px; font-size: 0.8rem; font-weight: 600; }
    .badge-warning { background: #FEF3C7; color: #D97706; padding: 3px 10px; border-radius: 12px; font-size: 0.8rem; font-weight: 600; }
    .badge-healthy { background: #DCFCE7; color: #16A34A; padding: 3px 10px; border-radius: 12px; font-size: 0.8rem; font-weight: 600; }
    .badge-info { background: #DBEAFE; color: #2563EB; padding: 3px 10px; border-radius: 12px; font-size: 0.8rem; font-weight: 600; }

    /* Section headers */
    .section-header {
        font-size: 1.3rem;
        font-weight: 600;
        color: #1E293B;
        border-bottom: 2px solid #2563EB;
        padding-bottom: 0.5rem;
        margin: 1.5rem 0 1rem 0;
    }

    /* Chat styling */
    .chat-user {
        background: #2563EB;
        color: white;
        padding: 0.8rem 1.2rem;
        border-radius: 12px 12px 4px 12px;
        margin: 0.5rem 0;
        max-width: 80%;
        margin-left: auto;
    }
    .chat-assistant {
        background: #F1F5F9;
        color: #1E293B;
        padding: 0.8rem 1.2rem;
        border-radius: 12px 12px 12px 4px;
        margin: 0.5rem 0;
        max-width: 80%;
    }

    /* Sidebar */
    [data-testid="stSidebar"] {
        background-color: #F8FAFC;
    }

    /* Info boxes */
    .info-box {
        background: #EFF6FF;
        border-left: 4px solid #2563EB;
        padding: 1rem;
        border-radius: 0 8px 8px 0;
        margin: 0.5rem 0;
    }
    .warning-box {
        background: #FFFBEB;
        border-left: 4px solid #F59E0B;
        padding: 1rem;
        border-radius: 0 8px 8px 0;
        margin: 0.5rem 0;
    }
    .danger-box {
        background: #FEF2F2;
        border-left: 4px solid #EF4444;
        padding: 1rem;
        border-radius: 0 8px 8px 0;
        margin: 0.5rem 0;
    }

    /* Hide streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}

    /* Tab styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        background-color: #F1F5F9;
        border-radius: 8px;
        padding: 8px 16px;
    }
    .stTabs [aria-selected="true"] {
        background-color: #2563EB !important;
        color: white !important;
    }
</style>
""",
    unsafe_allow_html=True,
)


# ── Session State Initialization ─────────────────────────────────
def init_session_state():
    defaults = {
        "data_loaded": False,
        "raw_data": None,
        "cleaned_data": None,
        "patterns": [],
        "near_misses": [],
        "schedule": [],
        "health_scores": [],
        "analysis": None,
        "anomaly_data": None,
        "failure_predictions": [],
        "chat_history": [],
        "notifications": [],
        "technician_config": {
            "name": "",
            "email": "",
            "phone": "",
            "alert_enabled": True,
        },
        "analysis_done": False,
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val


init_session_state()


# ── Notification System ─────────────────────────────────────────────
def send_notification(title: str, message: str, severity: str = "info"):
    """Add a notification to the session state."""
    from datetime import datetime
    notification = {
        "id": len(st.session_state.notifications) + 1,
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "title": title,
        "message": message,
        "severity": severity,
        "read": False,
    }
    st.session_state.notifications.append(notification)


def check_and_send_alerts():
    """Check data and send automatic alerts for critical conditions."""
    if not st.session_state.data_loaded:
        return
    
    patterns = st.session_state.patterns
    near_misses = st.session_state.near_misses
    health_scores = st.session_state.health_scores
    
    for p in patterns:
        if p.get("severity") == "critical":
            send_notification(
                f"🔴 CRITICAL: {p['equipment_name']}",
                f"Degradation detected in {p['signal']}. Confidence: {p['confidence']}%",
                "critical"
            )
    
    for nm in near_misses:
        if nm.get("status") == "approaching_critical":
            send_notification(
                f"⚠️ Near-Miss: {nm['equipment_name']}",
                f"{nm['parameter']} at {nm['failure_probability']}% failure probability",
                "warning"
            )
    
    for hs in health_scores:
        if hs.get("status") == "critical":
            send_notification(
                f"🚨 Health Alert: {hs['equipment_name']}",
                f"Health score: {hs['health_score']:.1f}% - Immediate attention required",
                "critical"
            )


def render_notifications_panel():
    """Render the notifications UI in sidebar."""
    st.markdown("---")
    st.markdown("### 🔔 Notifications")
    
    unread = [n for n in st.session_state.notifications if not n.get("read", False)]
    if unread:
        st.markdown(f"**{len(unread)} unread alert(s)**")
    else:
        st.markdown("No new notifications")
    
    if st.session_state.notifications:
        with st.expander("View All Notifications"):
            for n in reversed(st.session_state.notifications[-10:]):
                icon = "🔴" if n["severity"] == "critical" else "🟡" if n["severity"] == "warning" else "ℹ️"
                st.markdown(f"**{icon} {n['title']}**")
                st.caption(f"{n['timestamp']} - {n['message']}")
                if not n.get("read"):
                    if st.button("Mark Read", key=f"read_{n['id']}"):
                        n["read"] = True
                        st.rerun()
    
    if st.button("Clear All", key="clear_notifications"):
        st.session_state.notifications = []
        st.rerun()


# ── Sidebar ──────────────────────────────────────────────────────
def render_sidebar():
    with st.sidebar:
        st.markdown("## ⚙️ PredictiveAI")
        st.markdown("---")

        # Data upload
        st.markdown("### 📂 Data Upload")
        uploaded_file = st.file_uploader("Upload CSV", type=["csv"], help="Upload equipment log CSV file")
        if uploaded_file is not None:
            process_upload(uploaded_file)

        if st.button("📋 Load Sample Data", type="primary"):
            load_sample_data()

        if st.session_state.data_loaded:
            st.success(f"Data loaded: {len(st.session_state.cleaned_data)} rows")

        st.markdown("---")
        st.caption("PredictiveAI v2.0 | Streamlit Edition")


def process_upload(uploaded_file):
    """Process an uploaded CSV file."""
    file_id = getattr(uploaded_file, "file_id", uploaded_file.name)
    if st.session_state.get("current_file_id") == file_id:
        return
        
    try:
        uploaded_file.seek(0)
        content = uploaded_file.read()

        encodings = ["utf-8", "utf-16", "utf-16-le", "utf-16-be", "latin-1"]
        df = None
        last_error = None
        
        for encoding in encodings:
            try:
                df = pd.read_csv(io.StringIO(content.decode(encoding)))
                break
            except (UnicodeDecodeError, UnicodeError) as e:
                last_error = e
                continue
        
        if df is None:
            st.sidebar.error(f"Could not decode file. Tried: {', '.join(encodings)}. Last error: {last_error}")
            return

        if df.empty:
            st.sidebar.error("The uploaded file is empty.")
            return

        original_cols = list(df.columns)
        df.columns = [str(c).strip().lower().replace(" ", "_") for c in df.columns]
        
        required_cols = ["timestamp", "equipment_id"]
        missing = [c for c in required_cols if c not in df.columns]
        if missing:
            st.sidebar.error(f"Missing required columns: {missing}. Your data needs: timestamp, equipment_id. Found: {list(df.columns)}")
            return

        # Enforce numeric types
        numeric_cols = ["temperature", "vibration", "pressure", "rpm"]
        for col in numeric_cols:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors="coerce")

        # Parse timestamps and remove invalid rows
        df["timestamp"] = pd.to_datetime(df["timestamp"], errors="coerce")
        df = df.dropna(subset=["timestamp", "equipment_id"])

        if df.empty:
            st.sidebar.error("The uploaded file resulted in empty data after cleaning timestamps.")
            return

        # Ensure severity column exists and has bounded values
        if "severity" not in df.columns:
            df["severity"] = "info"
        else:
            df["severity"] = df["severity"].fillna("info").astype(str).str.lower()
            df["severity"] = df["severity"].apply(lambda x: x if x in ["info", "warning", "critical"] else "info")

        st.sidebar.info(f"Detected columns: {list(df.columns)}")
        
        st.session_state.raw_data = df.copy()
        st.session_state.cleaned_data = df.copy()
        st.session_state.data_loaded = True
        st.session_state.analysis_done = False

        st.session_state.current_file_id = file_id
        run_full_analysis()
    except Exception as e:
        st.sidebar.error(f"Error processing file: {e}")


def load_sample_data():
    """Load the sample dataset."""
    sample_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "dataset", "sample_logs.csv")
    if os.path.exists(sample_path):
        df = pd.read_csv(sample_path)
        st.session_state.raw_data = df.copy()
        st.session_state.cleaned_data = df.copy()
        st.session_state.data_loaded = True
        st.session_state.analysis_done = False
        run_full_analysis()
    else:
        st.sidebar.error("Sample data file not found.")


def run_full_analysis(force: bool = False):
    """Run all analysis modules on the cleaned data."""
    df = st.session_state.cleaned_data
    if df is None:
        return
    
    if st.session_state.get("analysis_done") and not force:
        return

    with st.spinner("Running pattern detection..."):
        st.session_state.patterns = detect_degradation_patterns(df)
    with st.spinner("Detecting near-miss events..."):
        st.session_state.near_misses = detect_near_misses(df)
    with st.spinner("Computing health scores..."):
        st.session_state.health_scores = get_equipment_health_scores(df)
    with st.spinner("Generating maintenance schedule..."):
        st.session_state.schedule = generate_maintenance_schedule(
            df,
            st.session_state.patterns,
            st.session_state.near_misses,
            {},
        )

    # ML Features
    with st.spinner("Running anomaly detection..."):
        st.session_state.anomaly_data = detect_anomalies_isolation_forest(df)
    with st.spinner("Predicting failure probabilities..."):
        st.session_state.failure_predictions = predict_failure_probability(df)

    # AI Analysis
    logs_summary = {
        "total_records": len(df),
        "date_range": {
            "start": str(df["timestamp"].min()) if "timestamp" in df.columns else None,
            "end": str(df["timestamp"].max()) if "timestamp" in df.columns else None,
        },
        "equipment_count": df["equipment_id"].nunique() if "equipment_id" in df.columns else 0,
        "severity_distribution": df["severity"].value_counts().to_dict() if "severity" in df.columns else {},
    }

    # AI Analysis (with timeout handling)
    try:
        with st.spinner("Running AI analysis (this may take a moment)..."):
            st.session_state.analysis = analyze_logs_with_llm(
                logs_summary,
                st.session_state.patterns,
                st.session_state.near_misses,
            )
    except Exception as e:
        st.session_state.analysis = {
            "error": str(e),
            "executive_summary": "AI analysis temporarily unavailable. Rule-based analysis is active.",
            "critical_findings": [],
            "root_causes": [],
            "immediate_actions": [],
            "long_term_strategy": []
        }
    
    st.session_state.analysis_done = True


# ── Main Content ─────────────────────────────────────────────────
def render_header():
    st.markdown(
        """
    <div class="main-header">
        <h1>⚙️ PredictiveAI - Maintenance Intelligence</h1>
        <p>AI-powered predictive maintenance scheduling for manufacturing equipment</p>
    </div>
    """,
        unsafe_allow_html=True,
    )


def render_landing():
    """Render landing page when no data is loaded."""
    st.markdown("### Welcome to PredictiveAI")
    st.markdown(
        "Upload your equipment log CSV or load the sample dataset from the sidebar to get started."
    )

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown(
            """
        <div class="metric-card">
            <div class="metric-value">5</div>
            <div class="metric-label">LLM Providers Supported</div>
        </div>
        """,
            unsafe_allow_html=True,
        )
    with col2:
        st.markdown(
            """
        <div class="metric-card">
            <div class="metric-value">9+</div>
            <div class="metric-label">Analysis Features</div>
        </div>
        """,
            unsafe_allow_html=True,
        )
    with col3:
        st.markdown(
            """
        <div class="metric-card">
            <div class="metric-value">ML</div>
            <div class="metric-label">Anomaly & Failure Prediction</div>
        </div>
        """,
            unsafe_allow_html=True,
        )

    st.markdown("---")
    st.markdown("#### Features")
    features = {
        "📈 Degradation Pattern Detection": "Identify temperature, vibration, and pressure anomalies",
        "⚠️ Near-Miss Detection": "Spot parameters approaching critical thresholds",
        "📅 Confidence-Based Scheduling": "Maintenance schedules with confidence scores",
        "🤖 Multi-LLM Chat Assistant": "Chat with 5 different AI providers",
        "🔬 ML Anomaly Detection": "Isolation Forest-based anomaly detection",
        "📊 Failure Prediction": "ML-based failure probability estimation",
        "📄 Report Export": "Download comprehensive maintenance reports",
    }
    cols = st.columns(3)
    for i, (title, desc) in enumerate(features.items()):
        with cols[i % 3]:
            st.markdown(f"**{title}**")
            st.caption(desc)


@st.fragment
def render_dashboard():
    """Render main dashboard with all tabs."""
    df = st.session_state.cleaned_data
    health_scores = st.session_state.health_scores

    # Summary metrics
    col1, col2, col3, col4, col5 = st.columns(5)
    critical_count = len([h for h in health_scores if h["status"] == "critical"])
    warning_count = len([h for h in health_scores if h["status"] in ["warning", "attention"]])
    healthy_count = len([h for h in health_scores if h["status"] == "healthy"])

    with col1:
        st.metric("Total Equipment", df["equipment_id"].nunique() if "equipment_id" in df.columns else 0)
    with col2:
        st.metric("Total Records", len(df))
    with col3:
        st.metric("Critical", critical_count, delta=None)
    with col4:
        st.metric("Warning", warning_count, delta=None)
    with col5:
        st.metric("Healthy", healthy_count, delta=None)

    st.markdown("---")

    # Main tabs
    tabs = st.tabs([
        "📊 Overview",
        "🔍 Patterns",
        "⚠️ Near-Misses",
        "📅 Schedule",
        "🔬 Anomaly Detection",
        "📈 Failure Prediction",
        "🤖 AI Chat",
        "📊 Data Explorer",
        "📄 Export Report",
    ])

    with tabs[0]:
        render_overview_tab()
    with tabs[1]:
        render_patterns_tab()
    with tabs[2]:
        render_near_misses_tab()
    with tabs[3]:
        render_schedule_tab()
    with tabs[4]:
        render_anomaly_tab()
    with tabs[5]:
        render_failure_prediction_tab()
    with tabs[6]:
        render_chat_tab()
    with tabs[7]:
        render_data_explorer_tab()
    with tabs[8]:
        render_export_tab()


# ── Tab Renderers ────────────────────────────────────────────────
def render_overview_tab():
    df = st.session_state.cleaned_data
    health_scores = st.session_state.health_scores
    analysis = st.session_state.analysis

    if df is None or "timestamp" not in df.columns:
        st.warning("No data with timestamp column available.")
        return

    # AI Analysis Summary
    if analysis and analysis.get("success"):
        a = analysis.get("analysis", {})
        source = analysis.get("source", "unknown")
        st.markdown(f'<div class="section-header">AI Analysis <span class="badge-info">{source.upper()}</span></div>', unsafe_allow_html=True)

        if isinstance(a, dict):
            if "executive_summary" in a:
                st.info(a["executive_summary"])
            elif "raw_response" in a:
                st.info(a["raw_response"][:500])

            # Critical findings
            if "critical_findings" in a and a["critical_findings"]:
                st.markdown("**Critical Findings:**")
                for f in a["critical_findings"]:
                    if isinstance(f, dict):
                        risk = f.get("risk_level", "MEDIUM")
                        badge_class = "badge-critical" if risk == "HIGH" else "badge-warning"
                        st.markdown(
                            f'<div class="warning-box"><span class="{badge_class}">{risk}</span> '
                            f'<strong>{f.get("equipment", "")}</strong>: {f.get("finding", "")}</div>',
                            unsafe_allow_html=True,
                        )

    st.markdown("---")

    # Equipment Health Overview
    st.markdown('<div class="section-header">Equipment Health Overview</div>', unsafe_allow_html=True)
    if health_scores:
        # Health score bar chart
        health_df = pd.DataFrame(health_scores)
        fig = px.bar(
            health_df,
            x="equipment_name",
            y="health_score",
            color="status",
            color_discrete_map={
                "healthy": "#22C55E",
                "attention": "#F59E0B",
                "warning": "#EF4444",
                "critical": "#DC2626",
            },
            title="Equipment Health Scores",
            labels={"equipment_name": "Equipment", "health_score": "Health Score"},
        )
        fig.update_layout(
            plot_bgcolor="white",
            paper_bgcolor="white",
            font_color="#1E293B",
            xaxis_tickangle=-45,
        )
        fig.add_hline(y=60, line_dash="dash", line_color="orange", annotation_text="Warning threshold")
        fig.add_hline(y=40, line_dash="dash", line_color="red", annotation_text="Critical threshold")
        st.plotly_chart(fig, width='stretch', key="overview_health")

    # Sensor trend charts
    st.markdown('<div class="section-header">Sensor Trends</div>', unsafe_allow_html=True)
    trend_data = get_trend_data(df)

    col1, col2 = st.columns(2)
    with col1:
        if "temperature" in trend_data.get("series", {}):
            fig = go.Figure()
            fig.add_trace(go.Scatter(
                x=trend_data["timestamps"],
                y=trend_data["series"]["temperature"],
                mode="lines+markers",
                name="Temperature (F)",
                line=dict(color="#EF4444"),
            ))
            fig.update_layout(title="Temperature Trends", plot_bgcolor="white", paper_bgcolor="white", font_color="#1E293B")
            st.plotly_chart(fig, width='stretch', key="temp_trend")

    with col2:
        if "vibration" in trend_data.get("series", {}):
            fig = go.Figure()
            fig.add_trace(go.Scatter(
                x=trend_data["timestamps"],
                y=trend_data["series"]["vibration"],
                mode="lines+markers",
                name="Vibration (mm/s)",
                line=dict(color="#F59E0B"),
            ))
            fig.update_layout(title="Vibration Trends", plot_bgcolor="white", paper_bgcolor="white", font_color="#1E293B")
            st.plotly_chart(fig, width='stretch', key="vibration_trend")

    col3, col4 = st.columns(2)
    with col3:
        if "pressure" in trend_data.get("series", {}):
            fig = go.Figure()
            fig.add_trace(go.Scatter(
                x=trend_data["timestamps"],
                y=trend_data["series"]["pressure"],
                mode="lines+markers",
                name="Pressure (PSI)",
                line=dict(color="#2563EB"),
            ))
            fig.update_layout(title="Pressure Trends", plot_bgcolor="white", paper_bgcolor="white", font_color="#1E293B")
            st.plotly_chart(fig, width='stretch', key="pressure_trend")

    with col4:
        # Severity distribution pie chart
        if "severity" in df.columns:
            sev_counts = df["severity"].value_counts()
            fig = px.pie(
                names=sev_counts.index,
                values=sev_counts.values,
                title="Severity Distribution",
                color=sev_counts.index,
                color_discrete_map={"info": "#3B82F6", "warning": "#F59E0B", "critical": "#EF4444"},
            )
            fig.update_layout(plot_bgcolor="white", paper_bgcolor="white", font_color="#1E293B")
            st.plotly_chart(fig, width='stretch', key="severity_dist")


def render_patterns_tab():
    patterns = st.session_state.patterns
    st.markdown(f'<div class="section-header">Degradation Patterns ({len(patterns)} detected)</div>', unsafe_allow_html=True)

    if not patterns:
        st.info("No degradation patterns detected. Upload data to analyze.")
        return

    for i, p in enumerate(patterns):
        severity = p["severity"]
        badge_class = "badge-critical" if severity == "critical" else "badge-warning"

        with st.expander(f"{'🔴' if severity == 'critical' else '🟡'} {p['equipment_name']} - {p['signal']} (Confidence: {p['confidence']}%)"):
            st.markdown(f'<span class="{badge_class}">{severity.upper()}</span>', unsafe_allow_html=True)
            st.markdown(f"**Description:** {p['description']}")
            st.markdown(f"**Recommendation:** {p['recommendation']}")

            # Confidence gauge
            fig = go.Figure(go.Indicator(
                mode="gauge+number",
                value=p["confidence"],
                title={"text": "Confidence"},
                gauge={
                    "axis": {"range": [0, 100]},
                    "bar": {"color": "#EF4444" if p["confidence"] > 70 else "#F59E0B"},
                    "steps": [
                        {"range": [0, 40], "color": "#DCFCE7"},
                        {"range": [40, 70], "color": "#FEF3C7"},
                        {"range": [70, 100], "color": "#FEE2E2"},
                    ],
                },
            ))
            fig.update_layout(height=250, plot_bgcolor="white", paper_bgcolor="white", font_color="#1E293B")
            st.plotly_chart(fig, width='stretch', key=f"pattern_gauge_{i}")

            st.markdown("**Evidence:**")
            for e in p.get("evidence", []):
                st.markdown(f"- {e}")


def render_near_misses_tab():
    near_misses = st.session_state.near_misses
    st.markdown(f'<div class="section-header">Near-Miss Events ({len(near_misses)} detected)</div>', unsafe_allow_html=True)

    if not near_misses:
        st.info("No near-miss events detected.")
        return

    for i, nm in enumerate(near_misses):
        status_icon = "🔴" if nm["status"] == "approaching_critical" else "🟡"
        with st.expander(f"{status_icon} {nm['equipment_name']} - {nm['parameter']} (Failure Prob: {nm['failure_probability']}%)"):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Current Value", f"{nm['current_value']:.1f} {nm['unit']}")
            with col2:
                st.metric("Warning Threshold", f"{nm['warning_threshold']} {nm['unit']}")
            with col3:
                st.metric("Critical Threshold", f"{nm['critical_threshold']} {nm['unit']}")

            # Threshold visualization
            fig = go.Figure()
            fig.add_trace(go.Bar(
                x=[nm["current_value"]],
                y=[nm["equipment_name"]],
                orientation="h",
                name="Current",
                marker_color="#F59E0B" if nm["status"] != "approaching_critical" else "#EF4444",
            ))
            fig.add_vline(x=nm["warning_threshold"], line_dash="dash", line_color="orange", annotation_text="Warning")
            fig.add_vline(x=nm["critical_threshold"], line_dash="dash", line_color="red", annotation_text="Critical")
            fig.update_layout(
                height=150,
                showlegend=False,
                plot_bgcolor="white",
                paper_bgcolor="white",
                font_color="#1E293B",
                title=f"{nm['parameter']} Level",
            )
            st.plotly_chart(fig, width='stretch', key=f"near_miss_chart_{i}")

            st.markdown(f"**Description:** {nm['description']}")
            st.markdown(f"**Action Required:** {nm['action_required']}")
            st.markdown(f"**Occurrences:** {nm['occurrences']} | **Last:** {nm['last_occurrence']}")


def render_schedule_tab():
    schedule = st.session_state.schedule
    st.markdown(f'<div class="section-header">Maintenance Schedule ({len(schedule)} tasks)</div>', unsafe_allow_html=True)

    if not schedule:
        st.info("No maintenance tasks scheduled. Upload data to generate a schedule.")
        return

    # Summary by urgency
    urgency_counts = {}
    for s in schedule:
        u = s["urgency"]
        urgency_counts[u] = urgency_counts.get(u, 0) + 1

    cols = st.columns(len(urgency_counts))
    color_map = {"IMMEDIATE": "#EF4444", "HIGH": "#F97316", "MEDIUM": "#F59E0B", "LOW": "#22C55E"}
    for i, (urgency, count) in enumerate(urgency_counts.items()):
        with cols[i]:
            st.markdown(
                f'<div class="metric-card"><div class="metric-value" style="color:{color_map.get(urgency, "#64748B")}">{count}</div>'
                f'<div class="metric-label">{urgency}</div></div>',
                unsafe_allow_html=True,
            )

    st.markdown("---")

    # Timeline chart
    schedule_df = pd.DataFrame(schedule)
    if "recommended_date" in schedule_df.columns:
        fig = px.scatter(
            schedule_df,
            x="recommended_date",
            y="equipment_name",
            color="urgency",
            size="confidence",
            hover_data=["task", "estimated_duration"],
            title="Maintenance Timeline",
            color_discrete_map=color_map,
        )
        fig.update_layout(plot_bgcolor="white", paper_bgcolor="white", font_color="#1E293B")
        st.plotly_chart(fig, width='stretch', key="maintenance_timeline")

    # Detailed list
    for s in schedule:
        urgency = s["urgency"]
        icon = "🔴" if urgency == "IMMEDIATE" else "🟠" if urgency == "HIGH" else "🟡" if urgency == "MEDIUM" else "🟢"
        with st.expander(f"{icon} [{urgency}] {s['equipment_name']}: {s['task']}"):
            st.markdown(f"**Description:** {s['description']}")
            st.markdown(f"**Confidence:** {s['confidence']}%")
            st.markdown(f"**Recommended Date:** {s['recommended_date']}")
            st.markdown(f"**Duration:** {s['estimated_duration']}")
            st.markdown(f"**Reasoning:** {s['reasoning']}")
            st.markdown(f"**Production Impact:** {s['production_impact']}")
            if s.get("evidence"):
                st.markdown("**Evidence:**")
                for e in s["evidence"]:
                    st.markdown(f"- {e}")




def render_anomaly_tab():
    st.markdown('<div class="section-header">ML Anomaly Detection (Isolation Forest)</div>', unsafe_allow_html=True)

    anomaly_data = st.session_state.anomaly_data
    if anomaly_data is None or anomaly_data.empty:
        st.info("No data available for anomaly detection. Upload data first.")
        return

    anomaly_count = len(anomaly_data[anomaly_data["anomaly"] == -1])
    normal_count = len(anomaly_data[anomaly_data["anomaly"] == 1])

    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Records", len(anomaly_data))
    with col2:
        st.metric("Anomalies Detected", anomaly_count)
    with col3:
        st.metric("Normal Records", normal_count)

    # Anomaly scatter plot
    if "temperature" in anomaly_data.columns and "vibration" in anomaly_data.columns:
        fig = px.scatter(
            anomaly_data,
            x="temperature",
            y="vibration",
            color=anomaly_data["anomaly"].map({1: "Normal", -1: "Anomaly"}),
            color_discrete_map={"Normal": "#22C55E", "Anomaly": "#EF4444"},
            title="Anomaly Detection: Temperature vs Vibration",
            labels={"color": "Classification"},
            hover_data=["equipment_name", "severity"] if "equipment_name" in anomaly_data.columns else None,
        )
        fig.update_layout(plot_bgcolor="white", paper_bgcolor="white", font_color="#1E293B")
        st.plotly_chart(fig, width='stretch', key="anomaly_temp_vib")

    if "pressure" in anomaly_data.columns and "temperature" in anomaly_data.columns:
        fig = px.scatter(
            anomaly_data,
            x="temperature",
            y="pressure",
            color=anomaly_data["anomaly"].map({1: "Normal", -1: "Anomaly"}),
            color_discrete_map={"Normal": "#22C55E", "Anomaly": "#EF4444"},
            title="Anomaly Detection: Temperature vs Pressure",
            labels={"color": "Classification"},
            hover_data=["equipment_name", "severity"] if "equipment_name" in anomaly_data.columns else None,
        )
        fig.update_layout(plot_bgcolor="white", paper_bgcolor="white", font_color="#1E293B")
        st.plotly_chart(fig, width='stretch', key="anomaly_temp_press")

    # Anomaly score distribution
    if "anomaly_score" in anomaly_data.columns:
        fig = px.histogram(
            anomaly_data,
            x="anomaly_score",
            nbins=30,
            title="Anomaly Score Distribution",
            color=anomaly_data["anomaly"].map({1: "Normal", -1: "Anomaly"}),
            color_discrete_map={"Normal": "#22C55E", "Anomaly": "#EF4444"},
        )
        fig.update_layout(plot_bgcolor="white", paper_bgcolor="white", font_color="#1E293B")
        st.plotly_chart(fig, width='stretch', key="anomaly_score_dist")

    # Anomalous records table
    anomalies = anomaly_data[anomaly_data["anomaly"] == -1]
    if not anomalies.empty:
        st.markdown("**Anomalous Records:**")
        display_cols = [c for c in ["timestamp", "equipment_name", "severity", "temperature", "vibration", "pressure", "message", "anomaly_score"] if c in anomalies.columns]
        st.dataframe(anomalies[display_cols].sort_values("anomaly_score"), width='stretch')


def render_failure_prediction_tab():
    st.markdown('<div class="section-header">ML Failure Prediction</div>', unsafe_allow_html=True)

    predictions = st.session_state.failure_predictions
    if not predictions:
        st.info("No predictions available. Upload data first.")
        return

    # Risk overview
    pred_df = pd.DataFrame(predictions)
    fig = px.bar(
        pred_df,
        x="equipment_name",
        y="failure_probability",
        color="risk_level",
        color_discrete_map={
            "CRITICAL": "#DC2626",
            "HIGH": "#F97316",
            "MEDIUM": "#F59E0B",
            "LOW": "#22C55E",
        },
        title="Failure Probability by Equipment",
        labels={"equipment_name": "Equipment", "failure_probability": "Failure Probability (%)"},
    )
    fig.update_layout(plot_bgcolor="white", paper_bgcolor="white", font_color="#1E293B", xaxis_tickangle=-45)
    fig.add_hline(y=70, line_dash="dash", line_color="red", annotation_text="Critical")
    fig.add_hline(y=50, line_dash="dash", line_color="orange", annotation_text="High")
    st.plotly_chart(fig, width='stretch', key="failure_prob_chart")

    # Detailed cards
    for p in predictions:
        risk = p["risk_level"]
        icon = "🔴" if risk == "CRITICAL" else "🟠" if risk == "HIGH" else "🟡" if risk == "MEDIUM" else "🟢"
        with st.expander(f"{icon} {p['equipment_name']} - {p['failure_probability']}% Failure Probability"):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Risk Level", risk)
            with col2:
                st.metric("Failure Probability", f"{p['failure_probability']}%")
            with col3:
                st.metric("Est. Time to Failure", p["estimated_time_to_failure"])

            st.markdown(f"**Recommended Action:** {p['recommended_action']}")
            st.markdown("**Contributing Factors:**")
            for f in p["contributing_factors"]:
                st.markdown(f"- {f}")




def render_chat_tab():
    st.markdown('<div class="section-header">AI Chat Assistant</div>', unsafe_allow_html=True)
    st.caption("Ask questions about your equipment, maintenance schedules, and more.")

    # Display chat history
    for msg in st.session_state.chat_history:
        if msg["role"] == "user":
            st.markdown(f'<div class="chat-user">{msg["content"]}</div>', unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="chat-assistant">{msg["content"]}</div>', unsafe_allow_html=True)

    # Chat input using a form to avoid StreamlitAPIException in tabs
    with st.form("chat_form", clear_on_submit=True):
        col1, col2 = st.columns([5, 1])
        with col1:
            user_input = st.text_input("Ask about maintenance, equipment status, schedules...", label_visibility="collapsed", placeholder="Type your message here...")
        with col2:
            submit = st.form_submit_button("Send")
            
    if submit and user_input:
        st.session_state.chat_history.append({"role": "user", "content": user_input})

        context = {
            "equipment_health": st.session_state.get("health_scores", []),
            "patterns": st.session_state.get("patterns", []),
            "schedule": st.session_state.get("schedule", []),
            "near_misses": st.session_state.get("near_misses", []),
        }

        result = chat_query(
            user_input,
            context,
            st.session_state.chat_history,
        )

        response = result.get("response", "I couldn't generate a response.")
        st.session_state.chat_history.append({"role": "assistant", "content": response})
        st.rerun()

    if st.button("Clear Chat History"):
        st.session_state.chat_history = []
        st.rerun()




def render_data_explorer_tab():
    st.markdown('<div class="section-header">Data Explorer</div>', unsafe_allow_html=True)

    df = st.session_state.cleaned_data
    if df is None or df.empty:
        st.info("No data available.")
        return

    # Filters
    col1, col2, col3 = st.columns(3)
    with col1:
        eq_filter = st.multiselect("Equipment", df["equipment_id"].unique().tolist() if "equipment_id" in df.columns else [])
    with col2:
        sev_filter = st.multiselect("Severity", df["severity"].unique().tolist() if "severity" in df.columns else [])
    with col3:
        type_filter = st.multiselect("Log Type", df["log_type"].unique().tolist() if "log_type" in df.columns else [])

    filtered = df.copy()
    if eq_filter:
        filtered = filtered[filtered["equipment_id"].isin(eq_filter)]
    if sev_filter:
        filtered = filtered[filtered["severity"].isin(sev_filter)]
    if type_filter:
        filtered = filtered[filtered["log_type"].isin(type_filter)]

    st.markdown(f"Showing **{len(filtered)}** of **{len(df)}** records")
    st.dataframe(filtered, width='stretch', height=500)

    # Download filtered data
    from io import BytesIO
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Filtered Equipment Data', 0)
        doc.add_paragraph(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Records: {len(filtered)}")
        
        # Add table with data
        table = doc.add_table(rows=1, cols=len(filtered.columns))
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        for i, col in enumerate(filtered.columns):
            header_cells[i].text = str(col)
        
        # Data rows (limit to first 100 for performance)
        for _, row in filtered.head(100).iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
        
        if len(filtered) > 100:
            doc.add_paragraph(f"... and {len(filtered) - 100} more rows")
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            "📄 Download Data (DOCX)",
            buffer.getvalue(),
            f"filtered_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except ImportError:
        csv = filtered.to_csv(index=False)
        st.download_button("Download Filtered Data (CSV)", csv, "filtered_data.csv", "text/csv")


def render_export_tab():
    st.markdown('<div class="section-header">Export Report</div>', unsafe_allow_html=True)

    if not st.session_state.data_loaded:
        st.info("Upload data first to generate a report.")
        return

    st.markdown("Generate a comprehensive maintenance report with all analysis results.")

    report = generate_report_data(
        st.session_state.cleaned_data,
        st.session_state.patterns,
        st.session_state.near_misses,
        st.session_state.schedule,
        st.session_state.health_scores,
        st.session_state.analysis,
    )

    st.text_area("Report Preview", report, height=400)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button(
            "📄 Download Report (TXT)",
            report,
            f"maintenance_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
            "text/plain",
            type="primary",
        )
    with col2:
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document()
            doc.add_heading('PredictiveAI Maintenance Report', 0)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_heading('Analysis Summary', level=1)
            
            if st.session_state.analysis:
                doc.add_paragraph(str(st.session_state.analysis))
            
            doc.add_heading('Patterns Detected', level=1)
            for p in st.session_state.patterns:
                doc.add_paragraph(f"- {p}")
            
            doc.add_heading('Near Misses', level=1)
            for nm in st.session_state.near_misses:
                doc.add_paragraph(f"- {nm}")
            
            doc.add_heading('Maintenance Schedule', level=1)
            for s in st.session_state.schedule:
                doc.add_paragraph(f"- {s}")
            
            doc.add_heading('Equipment Health Scores', level=1)
            for hs in st.session_state.health_scores:
                doc.add_paragraph(f"{hs['equipment_name']}: {hs['health_score']}")
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.download_button(
                "📄 Download Report (DOCX)",
                buffer.getvalue(),
                f"maintenance_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
            )
        except ImportError:
            st.warning("DOCX export not available. Install python-docx package.")
    
    with col3:
        # CSV export of all analysis data
        export_data = {
            "patterns": st.session_state.patterns,
            "near_misses": st.session_state.near_misses,
            "schedule": st.session_state.schedule,
            "health_scores": st.session_state.health_scores,
        }
        import json
        json_export = json.dumps(export_data, indent=2, default=str)
        st.download_button(
            "📊 Download Analysis Data (JSON)",
            json_export,
            f"analysis_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            "application/json",
            type="primary",
        )


# ── Main ─────────────────────────────────────────────────────────
def main():
    render_sidebar()
    render_header()

    if st.session_state.data_loaded:
        render_dashboard()
    else:
        render_landing()


if __name__ == "__main__":
    main()
